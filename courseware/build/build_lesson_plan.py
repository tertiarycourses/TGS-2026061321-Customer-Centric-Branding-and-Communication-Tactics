#!/usr/bin/env python3
"""Generate the Customer-Centric Branding and Communication Tactics Lesson
Plan (LP) DOCX in the Tertiary WSQ house format.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 2-day schedule table (Day 1 9:30am-6:30pm full teaching day;
Day 2 9:30am-4:00pm teaching then the 4:00-6:00pm Assessment block).
The Daily Schedule table's "Slides" column is read from slide_map.json,
written by build_slides.py, so the LP always cites the actual deck page —
ALWAYS run build_slides.py before this script.
"""
import os, sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4
import prodoc
REPO=os.path.dirname(os.path.dirname(HERE)); ASSETS=os.path.join(REPO,"courseware","assets")

with open(os.path.join(HERE,"slide_map.json")) as f:
    SLIDE_MAP=json.load(f)

def act_title(n): return next(a["title"] for a in ACT if a["num"]==n)
def slides_for(*keys):
    pages=[SLIDE_MAP[k] for k in keys if SLIDE_MAP.get(k)]
    if not pages: return "—"
    lo,hi=min(pages),max(pages)
    return f"Slide {lo}" if lo==hi else f"Slides {lo}–{hi}"

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"; ADMIN_FILL="F3F5F8"

# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, text, slide_keys)  kind: admin/topic/activity/break/lunch/assess/recap
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:30","9:50",20,"admin","Welcome, course introduction, ground rules, Skills Framework overview and mandatory digital attendance (AM)",[]),
    ("9:50","11:30",100,"topic",f"LU1 — Stakeholders and Organisation: concepts, then {act_title(1)} & {act_title(2)}",["topic1_section","activity1","activity2"]),
    ("11:30","11:40",10,"break","Tea break",[]),
    ("11:40","12:30",50,"activity",f"{act_title(3)} & {act_title(4)}; LU1 recap",["activity3","activity4"]),
    ("12:30","1:30",60,"lunch","Lunch break",[]),
    ("1:30","3:10",100,"topic",f"LU2 — Customer Influence: concepts, then {act_title(5)} & {act_title(6)}",["topic2_section","activity5","activity6"]),
    ("3:10","3:20",10,"break","Tea break",[]),
    ("3:20","4:40",80,"activity",f"{act_title(7)} & {act_title(8)}; LU2 recap",["activity7","activity8"]),
    ("4:40","6:20",100,"topic",f"LU3 — Branding in Marketing (begins): concepts, then {act_title(9)} & {act_title(10)}",["topic3_section","activity9","activity10"]),
    ("6:20","6:30",10,"recap","Day 1 recap, Q&A and PM digital attendance",[]),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:30","9:40",10,"recap","Day 1 recap and mandatory digital attendance (AM)",[]),
    ("9:40","11:00",80,"activity",f"LU3 (cont'd): {act_title(11)} & {act_title(12)}; LU3 recap",["activity11","activity12"]),
    ("11:00","11:10",10,"break","Tea break",[]),
    ("11:10","12:30",80,"topic",f"LU4 — Branding Effectiveness: concepts, then {act_title(13)} & {act_title(14)}",["topic4_section","activity13","activity14"]),
    ("12:30","1:15",45,"lunch","Lunch break",[]),
    ("1:15","2:35",80,"activity",f"{act_title(15)} & {act_title(16)}",["activity15","activity16"]),
    ("2:35","2:45",10,"break","Tea break",[]),
    ("2:45","3:50",65,"activity",f"{act_title(17)}; LU4 recap; What You Achieved; Courseware & Assessment on the LMS",["activity17"]),
    ("3:50","4:00",10,"assess","Briefing for Assessment",["briefing_assessment"]),
    ("4:00","5:00",60,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, open book",["assessment_front"]),
    ("5:00","6:00",60,"assess","Case Study (CS) — 1 hour, open book. PM digital attendance and TRAQOM survey",["assessment_end"]),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.DOC_VERSION,
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[(C.DOC_VERSION,C.VERSION_DATE,"First version.",C.ORG)])
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("TSC Reference",f"{C.TSC_TITLE} ({C.TSC_CODE}) · {C.TSC_LEVEL}"),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration","2 days · 9:30am–6:30pm (Day 1) / 9:30am–6:00pm (Day 2)"),
      ("Daily Timing","1-hour lunch on Day 1, 45-minute lunch on Day 2; tea breaks within training time"),
      ("Mode","Instructor-led, with case-study and workshop activities per Learning Unit"),
      ("Instructional Methods","Interactive presentation, discussions, case studies, peer teaching / peer practice"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "The assessment block runs on Day 2 from 4:00pm to 6:00pm.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":ADMIN_FILL,"recap":ADMIN_FILL,"activity":None}

H("Daily Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=4); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity","Slides"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text,skeys in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[3],slides_for(*skeys),size=9.5,fill=fill,align=WD_ALIGN_PARAGRAPH.CENTER)
        if kind!="lunch": training+=mins
    for row in tbl.rows:
        row.cells[0].width=Inches(1.0); row.cells[1].width=Inches(0.8)
        row.cells[2].width=Inches(4.2); row.cells[3].width=Inches(1.1)
    p=doc.add_paragraph(); r=p.add_run(f"Total training + assessment time: {training} minutes ({training/60:.2f} hours)."); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    if day==2:
        p2=doc.add_paragraph(); r2=p2.add_run(
            "Day 2 lunch is compressed to 45 minutes (from Day 1's 1 hour) so that content and wrap-up finish "
            "by 4:00pm, per the WSQ two-day assessment-day scheduling rule — the 2-hour Assessment block "
            "(WA + Case Study) then runs 4:00pm–6:00pm, ending the course at 6:00pm.")
        r2.italic=True; r2.font.size=Pt(9.5); r2.font.color.rgb=GREY

H("Topic-by-Topic Breakdown",1)
for tp in C.TOPICS:
    topic_key=f"topic{tp['num']}_section"
    H(f"{tp['code']} — {tp['title']} · {slides_for(topic_key)}",2)
    doc.add_paragraph(tp["subtitle"])
    p=doc.add_paragraph("Key concepts:"); p.runs[0].bold=True
    for c in tp["concepts"]:
        doc.add_paragraph(c,style="List Bullet")
    for a in [x for x in ACT if x["topic"]==tp["num"]]:
        act_key=f"activity{a['num']}"
        H(f"Activity {a['num']} — {a['title']} · {slides_for(act_key)}",3)
        doc.add_paragraph(f"Objective: {a['objective']}.")
        doc.add_paragraph(f"Scenario: {a['desc']}")
        doc.add_paragraph(f"Duration: {a['duration']}.")

H("Resources Required",1)
for r in ["Projector/screen and PA system for the trainer's slide deck.",
          "Whiteboard/flip chart and markers for group activity work.",
          "Printed or digital copies of the Learner Guide for every learner.",
          "Internet access for learners to research live brand examples during activities.",
          "Mobile phones (learners' own) for the SSG digital-attendance QR scans."]:
    doc.add_paragraph(r,style="List Bullet")

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Grading: Competent / Not Yet Competent.",
          "A minimum of 75% attendance (SSG Digital Attendance record) is required for funding eligibility.",
          "Learners submit their assessment answers and complete the TRAQOM survey on the LMS at "
          "https://lms-tms.tertiaryinfotech.com/."]:
    doc.add_paragraph(a,style="List Bullet")

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
