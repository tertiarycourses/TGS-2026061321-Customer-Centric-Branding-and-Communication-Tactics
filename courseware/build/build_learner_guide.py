#!/usr/bin/env python3
"""Generate the Customer-Centric Branding and Communication Tactics Learner
Guide as a DOCX (courseware/LG-*.docx) — DOCX + PDF only, no Markdown mirror
kept in the repo (wsq-learner-guide HARD RULE 1).

House format: cover page, Document Version Control Record, auto TOC, Arial
11pt body, one section per Learning Unit with concepts, then one subsection
per activity (Objective · Scenario · Step-by-step · Debrief), a Quick
Reference table, Support section and the assessment flow. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned
with the slide deck and Lesson Plan.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4
import prodoc
REPO=os.path.dirname(os.path.dirname(HERE)); ASSETS=os.path.join(REPO,"courseware","assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)

doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.DOC_VERSION,
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[(C.DOC_VERSION,C.VERSION_DATE,"First version.",C.ORG)])
prodoc.add_toc(doc)

def h3(text,color=BRAND):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=color
    return p

doc.add_heading("How to Use This Guide",level=1)
doc.add_paragraph(
    f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
    f"It supports the Skills Framework Technical Skill & Competency “{C.TSC_TITLE}” "
    f"({C.TSC_CODE}, {C.TSC_LEVEL}) across 4 Learning Units and 17 in-class activities.")
doc.add_paragraph(
    "Use this guide alongside the course slides during class, and again during the open-book "
    "assessment. Each Learning Unit section below lists the key concepts, followed by every "
    "activity you will complete in class with its scenario, step-by-step instructions and a "
    "debrief prompt to check your own work. This is a workshop-based course rather than a "
    "software course, so activities are illustrated with worked templates and examples on the "
    "slides rather than software screenshots.")
doc.add_paragraph("Before you start, you will need:")
for b in ["A notebook or digital document to capture your activity work.",
          "Internet access to research live brand examples where an activity calls for one.",
          "The course slides (downloaded from the LMS) for reference during activities and the assessment."]:
    doc.add_paragraph(b,style="List Bullet")

doc.add_heading("Course Learning Outcomes",level=1)
for lo in C.LEARNING_OUTCOMES:
    doc.add_paragraph(lo,style="List Bullet")

doc.add_heading("Skills Framework Reference",level=1)
for line in [f"TSC Title: {C.TSC_TITLE}", f"TSC Code: {C.TSC_CODE}", f"Proficiency Level: {C.TSC_LEVEL}"]:
    doc.add_paragraph(line,style="List Bullet")

for t in C.TOPICS:
    doc.add_heading(f"{t['code']} — {t['title']}",level=1)
    doc.add_paragraph(t["subtitle"])
    h3("Key concepts")
    for c in t["concepts"]:
        doc.add_paragraph(c,style="List Bullet")
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        doc.add_heading(f"Activity {a['num']} — {a['title']}",level=2)
        doc.add_paragraph(f"Objective: {a['objective']}.")
        h3("Scenario")
        doc.add_paragraph(a["desc"])
        h3("You'll produce")
        doc.add_paragraph(f"{a['build']}   (Duration: {a['duration']}.)")
        h3("Step-by-step")
        for i,(instr,_cmd) in enumerate(a["steps"],1):
            p=doc.add_paragraph(style="List Number"); p.add_run(instr)
        h3("Debrief it")
        p=doc.add_paragraph(); r=p.add_run("Check: "); r.bold=True; r.font.color.rgb=BRAND
        p.add_run(a["test"]).font.size=Pt(10.5)
        doc.add_paragraph("")

doc.add_heading("Quick Reference — Activities by Learning Unit",level=1)
tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"
hdr=tbl.add_row().cells
for i,htext in enumerate(["Learning Unit","Activity","Duration"]):
    hdr[i].text=""; r=hdr[i].paragraphs[0].add_run(htext); r.bold=True; r.font.size=Pt(9.5)
    prodoc._shade_cell(hdr[i],"1F6FEB")
    r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for t in C.TOPICS:
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        cells=tbl.add_row().cells
        cells[0].text=""; cells[0].paragraphs[0].add_run(t["code"]).font.size=Pt(9.5)
        cells[1].text=""; cells[1].paragraphs[0].add_run(f"{a['num']}. {a['title']}").font.size=Pt(9.5)
        cells[2].text=""; cells[2].paragraphs[0].add_run(a["duration"]).font.size=Pt(9.5)

doc.add_heading("Assessment",level=1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — this Learner Guide, the course slides and approved materials only.",
          "Grading: Competent / Not Yet Competent.",C.ASSESSMENT["note"]]:
    doc.add_paragraph(a,style="List Bullet")

doc.add_heading("Assessment Flow",level=1)
for i,step in enumerate(["TRAQOM — scan the TRAQOM QR code on the LMS and complete the survey.",
                          "Assessment Digital Attendance.",
                          "Assessment (Written Assessment + Case Study).",
                          "Submit the assessment answers on the LMS.",
                          "Sign the Assessment Summary Record."],1):
    doc.add_paragraph(f"{i}. {step}")
doc.add_paragraph("Courseware and the assessment are on the LMS: https://lms-tms.tertiaryinfotech.com/")

doc.add_heading("Glossary",level=1)
for term,defn in [
    ("Brand positioning","The distinct place a brand occupies in the customer's mind relative to competitors."),
    ("Stakeholder","Any internal or external party with an interest in, or influence over, the brand."),
    ("Active listening","Giving full attention, applying a structured framework, and paraphrasing without judgement."),
    ("Brand equity","The commercial value derived from consumer perception of the brand, not the product itself."),
    ("KPI","Key Performance Indicator — a specific, measurable metric used to judge progress against a goal."),
    ("AMEC framework","An industry framework for measuring and evaluating PR / communication campaign effectiveness."),
    ("NPS","Net Promoter Score — a customer-loyalty metric based on likelihood to recommend the brand."),
    ("TRAQOM","The SSG post-course survey completed via the LMS, distinct from digital attendance."),
]:
    p=doc.add_paragraph(style="List Bullet")
    r=p.add_run(term+" — "); r.bold=True; p.add_run(defn)

doc.add_heading("Support",level=1)
for line in ["Email: enquiry@tertiaryinfotech.com","Tel: +65 6100 0613","Website: www.tertiarycourses.com.sg",
             "LMS: https://lms-tms.tertiaryinfotech.com/"]:
    doc.add_paragraph(line,style="List Bullet")

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
