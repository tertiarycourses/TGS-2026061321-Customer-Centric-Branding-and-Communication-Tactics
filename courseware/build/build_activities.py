#!/usr/bin/env python3
"""Generate the activities/ folder — ONE FOLDER PER ACTIVITY, each holding the
house artefact set as DOCX **and** PDF, matching the Tertiary Infotech reference
activity packs (reference/Activities/NN - Title/A NN-*.docx|pdf):

    activities/
      README.md
      tools.md
      01 - Stakeholder Influence Mapping/
        A01-Facilitator-Guide-Stakeholder-Influence-Mapping.docx | .pdf
        A01-Learner-Worksheet-Stakeholder-Influence-Mapping.docx | .pdf
        A01-Checklist-Stakeholder-Influence-Mapping.docx         | .pdf
      ...

House format taken from the reference pack:
  * blue banner table (org name + course title · course code)
  * a 4-chip meta row: TYPE / DURATION / MAPS TO / TOPIC
  * Name+Date table on learner-facing sheets
  * Heading 2 sections, Arial 11pt body
  * Facilitator Guide carries the numbered run-sheet as a 3-column table
    (# | Instruction to learners | Facilitator note)
  * Checklist carries a criteria table (# | Criterion | Evidences | ✓)

Content is driven by course_data.py + data_domain1..4.py + data_activity_steps.py
so the packs stay aligned with the slide deck, Lesson Plan and Learner Guide.

HOUSE RULE: the detailed step-by-step lives HERE and in the Learner Guide —
never on the slides.
"""
import os, re, sys, shutil, subprocess, tempfile

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_activity_steps import STEPS
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4
REPO=os.path.dirname(os.path.dirname(HERE))
ACTIVITIES=os.path.join(REPO,"activities")

BRAND="1F6FEB"; INK="111827"; GREEN="10B981"; LIGHT="F5F8FC"; LINE="E2E8F0"
BRAND_RGB=RGBColor(0x1F,0x6F,0xEB); INK_RGB=RGBColor(0x11,0x18,0x27)
GREY_RGB=RGBColor(0x55,0x5B,0x66)
RULE="_"*95

def slug(t): return re.sub(r"[^a-z0-9]+","-",t.lower()).strip("-")
def title_slug(t): return re.sub(r"[^A-Za-z0-9]+","-",t).strip("-")

TOPIC_BY_NUM={t["num"]:t for t in C.TOPICS}
TYPE_LABEL={"role_play":"Role Play","case_study":"Case Study"}

def folder_for(a): return f"{a['num']:02d} - {a['title']}"

# ------------------------------------------------------------------ docx bits
def shade(cell,hexfill):
    tcPr=cell._tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hexfill)
    tcPr.append(shd)

def no_borders(table):
    tbl=table._tbl; tblPr=tbl.tblPr
    borders=OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        el=OxmlElement(f"w:{edge}"); el.set(qn("w:val"),"none"); el.set(qn("w:sz"),"0")
        borders.append(el)
    tblPr.append(borders)

def new_doc():
    doc=Document()
    n=doc.styles["Normal"]; n.font.name="Arial"; n.font.size=Pt(11)
    for st,sz,col in (("Heading 1",16,BRAND),("Heading 2",13,INK)):
        s=doc.styles[st]; s.font.name="Arial"; s.font.size=Pt(sz); s.font.bold=True
        s.font.color.rgb=RGBColor.from_string(col)
    for s in doc.sections:
        s.left_margin=s.right_margin=Cm(2.26)
        s.top_margin=s.bottom_margin=Cm(2.0)
    return doc

def banner(doc):
    """Blue org/course banner across the top."""
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.rows[0].cells[0]; shade(c,BRAND); no_borders(t)
    p=c.paragraphs[0]; p.space_after=Pt(0)
    r=p.add_run(C.ORG); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    p2=c.add_paragraph(); p2.space_before=Pt(0)
    r2=p2.add_run(f"{C.TITLE}  ·  {C.COURSE_CODE}")
    r2.font.size=Pt(8.5); r2.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    return t

def meta_chips(doc,a):
    """TYPE / DURATION / MAPS TO / TOPIC chip row."""
    t=TOPIC_BY_NUM[a["topic"]]
    chips=[("TYPE",TYPE_LABEL.get(a.get("case_type"),"Workshop")),
           ("DURATION",a["duration"]),
           ("MAPS TO",t["code"]),
           ("TOPIC",f"{t['code']} — {t['title']}")]
    tbl=doc.add_table(rows=1,cols=4); no_borders(tbl)
    for cell,(k,v) in zip(tbl.rows[0].cells,chips):
        shade(cell,LIGHT)
        p=cell.paragraphs[0]; p.space_after=Pt(0)
        r=p.add_run(k); r.bold=True; r.font.size=Pt(7.5); r.font.color.rgb=BRAND_RGB
        p2=cell.add_paragraph(); p2.space_before=Pt(0)
        r2=p2.add_run(v); r2.font.size=Pt(9); r2.font.color.rgb=INK_RGB
    doc.add_paragraph()
    return tbl

def name_date(doc):
    t=doc.add_table(rows=2,cols=2); t.style="Table Grid"
    for i,label in enumerate(("Name","Date")):
        c=t.rows[i].cells[0]; c.text=""
        r=c.paragraphs[0].add_run(label); r.bold=True; r.font.size=Pt(9.5)
        shade(c,LIGHT)
        t.rows[i].cells[1].text=""
    doc.add_paragraph()

def head(doc,a,kind):
    p=doc.add_paragraph(); p.space_after=Pt(2)
    r=p.add_run(f"ACTIVITY {a['num']} — {kind}")
    r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BRAND_RGB
    p2=doc.add_paragraph(); p2.space_after=Pt(8)
    r2=p2.add_run(a["title"]); r2.bold=True; r2.font.size=Pt(15); r2.font.color.rgb=INK_RGB

def blanks(doc,n=1):
    for _ in range(n):
        p=doc.add_paragraph(); p.space_after=Pt(2)
        r=p.add_run(RULE); r.font.size=Pt(10); r.font.color.rgb=RGBColor(0xAA,0xB2,0xBF)

def footer_note(doc):
    p=doc.add_paragraph(); p.space_before=Pt(14)
    r=p.add_run(f"{C.TITLE} · {C.COURSE_CODE} · Version {C.VERSION} · "
                f"© 2026 {C.ORG}")
    r.font.size=Pt(8); r.font.color.rgb=GREY_RGB; r.italic=True

# ------------------------------------------------------------------ documents
def facilitator_guide(a):
    sp=STEPS[a["num"]]; t=TOPIC_BY_NUM[a["topic"]]
    doc=new_doc(); banner(doc); head(doc,a,"FACILITATOR GUIDE"); meta_chips(doc,a)

    doc.add_heading("Purpose of this activity",level=2)
    doc.add_paragraph(f"By the end of this activity learners will be able to {a['objective'][0].lower()}"
                      f"{a['objective'][1:]} ({t['code']}).")

    doc.add_heading("What learners produce",level=2)
    doc.add_paragraph(sp["artefact"])

    doc.add_heading("Materials required",level=2)
    for m in ["Learner Worksheet (one per learner)","Checklist (one per group)",
              "Flipchart or A3 sheet and markers","The course slides for the scenario and framework",
              "Internet access where the activity calls for live research"]:
        doc.add_paragraph(m,style="List Bullet")

    doc.add_heading("Set-up",level=2)
    grouping=("Groups of exactly 3 so every role is played."
              if a.get("roles") else "Groups of 3–5.")
    doc.add_paragraph(f"{grouping} Distribute one worksheet per learner and one checklist per group. "
                      f"All 17 activities run on the same continuous Nimbus Wellness scenario, so remind "
                      f"learners of where this activity sits in that story before you start.")

    doc.add_heading("Scenario / briefing material",level=2)
    for para in a["case_scenario"]:
        doc.add_paragraph(para)

    if a.get("roles"):
        doc.add_heading("Roles",level=2)
        doc.add_paragraph("Assign one role per participant:")
        for name,goal,brief in a["roles"]:
            p=doc.add_paragraph(style="List Bullet")
            r=p.add_run(f"{name} — "); r.bold=True
            r2=p.add_run(f"Goal: {goal}. "); r2.italic=True
            p.add_run(brief)

    doc.add_heading("How to run it — step by step",level=2)
    tbl=doc.add_table(rows=1,cols=3); tbl.style="Table Grid"
    hdr=tbl.rows[0].cells
    for i,h in enumerate(["#","Instruction to learners","Facilitator note"]):
        hdr[i].text=""; r=hdr[i].paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(hdr[i],BRAND)
    tips=sp.get("tips",[])
    for i,(heading,instruction) in enumerate(sp["steps"],1):
        cells=tbl.add_row().cells
        cells[0].text=""; cells[0].paragraphs[0].add_run(str(i)).font.size=Pt(9.5)
        cells[1].text=""
        p=cells[1].paragraphs[0]
        r=p.add_run(f"{heading}. "); r.bold=True; r.font.size=Pt(9.5)
        r2=p.add_run(instruction); r2.font.size=Pt(9.5)
        cells[2].text=""
        note=tips[i-1] if i-1 < len(tips) else ""
        cells[2].paragraphs[0].add_run(note).font.size=Pt(9)
        if i%2==0:
            for c in cells: shade(c,LIGHT)

    doc.add_heading("Suggested timing",level=2)
    tt=doc.add_table(rows=1,cols=2); tt.style="Table Grid"
    h=tt.rows[0].cells
    for i,x in enumerate(["Minutes","Phase"]):
        h[i].text=""; r=h[i].paragraphs[0].add_run(x)
        r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(h[i],BRAND)
    for mins,phase in sp["timing"]:
        c=tt.add_row().cells
        c[0].text=""; c[0].paragraphs[0].add_run(f"{mins} min").font.size=Pt(9.5)
        c[1].text=""; c[1].paragraphs[0].add_run(phase).font.size=Pt(9.5)
    p=doc.add_paragraph()
    r=p.add_run(f"Total: {sum(m for m,_ in sp['timing'])} minutes"); r.bold=True; r.font.size=Pt(9.5)

    doc.add_heading("Discussion & decision prompts",level=2)
    for prompt in a["discussion_prompts"]:
        doc.add_paragraph(prompt,style="List Number")

    doc.add_heading("Debrief",level=2)
    doc.add_paragraph("Bring the class back together and draw out the learning:")
    for b in ["Ask two or three groups to show what they produced, not just what they concluded.",
              "Ask what surprised them — the gap between what they assumed and what they found.",
              f"Link the activity explicitly back to {t['code']} — {t['title']}.",
              "Remind learners this activity rehearses what the Case Study assessment will require."]:
        doc.add_paragraph(b,style="List Bullet")
    doc.add_paragraph("Reflection questions:")
    for point in a["reflection_points"]:
        doc.add_paragraph(point,style="List Bullet")

    if sp.get("tips"):
        doc.add_heading("Facilitator notes — what to watch for",level=2)
        for tip in sp["tips"]:
            doc.add_paragraph(tip,style="List Bullet")

    doc.add_heading("Success criterion",level=2)
    doc.add_paragraph(a["debrief_check"])

    doc.add_heading("Assessment link",level=2)
    doc.add_paragraph(
        f"This activity builds {t['code']} — {t['title']}. All 17 activities are in-class "
        f"activities assessed indirectly through the open-book Case Study (CS) component, which "
        f"reuses the same techniques on the same continuous Nimbus Wellness scenario. Knowledge "
        f"is assessed in the Written Assessment (WA-SAQ). Nothing is assessed that is not "
        f"practised here or taught in the slides.")
    footer_note(doc)
    return doc

def learner_worksheet(a):
    sp=STEPS[a["num"]]; t=TOPIC_BY_NUM[a["topic"]]
    doc=new_doc(); banner(doc); head(doc,a,"LEARNER WORKSHEET"); meta_chips(doc,a); name_date(doc)

    doc.add_heading("What you are doing",level=2)
    doc.add_paragraph(a["case_scenario"][0])
    p=doc.add_paragraph()
    r=p.add_run("What you will produce: "); r.bold=True
    p.add_run(sp["artefact"])

    doc.add_heading("The scenario",level=2)
    for para in a["case_scenario"][1:]:
        doc.add_paragraph(para)

    if a.get("roles"):
        doc.add_heading("Role assignment",level=2)
        rt=doc.add_table(rows=1,cols=3); rt.style="Table Grid"
        h=rt.rows[0].cells
        for i,x in enumerate(["Role","Played by","Goal"]):
            h[i].text=""; r=h[i].paragraphs[0].add_run(x)
            r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            shade(h[i],BRAND)
        for name,goal,brief in a["roles"]:
            c=rt.add_row().cells
            c[0].text=""; c[0].paragraphs[0].add_run(name).font.size=Pt(9.5)
            c[1].text=""
            c[2].text=""; c[2].paragraphs[0].add_run(goal).font.size=Pt(9.5)
        doc.add_paragraph()

    doc.add_heading("Your steps",level=2)
    for heading,instruction in sp["steps"]:
        p=doc.add_paragraph(style="List Number")
        r=p.add_run(f"{heading} — "); r.bold=True
        p.add_run(instruction)

    doc.add_heading("Your working space",level=2)
    doc.add_paragraph("Complete every field. You may keep this worksheet and refer to it during "
                      "the open-book assessment.")
    for i,(heading,_) in enumerate(sp["steps"],1):
        p=doc.add_paragraph(); p.space_after=Pt(2); p.space_before=Pt(8)
        r=p.add_run(f"Step {i} — {heading}"); r.bold=True; r.font.size=Pt(10)
        r.font.color.rgb=BRAND_RGB
        blanks(doc,2)

    doc.add_heading("Answers to the discussion prompts",level=2)
    for i,prompt in enumerate(a["discussion_prompts"],1):
        p=doc.add_paragraph(); p.space_after=Pt(2); p.space_before=Pt(8)
        r=p.add_run(f"{i}. {prompt}"); r.bold=True; r.font.size=Pt(10)
        blanks(doc,2)

    doc.add_heading("Reflect & discuss",level=2)
    for point in a["reflection_points"]:
        p=doc.add_paragraph(); p.space_after=Pt(2); p.space_before=Pt(8)
        r=p.add_run(point); r.bold=True; r.font.size=Pt(10)
        blanks(doc,2)

    doc.add_heading("Check your work",level=2)
    doc.add_paragraph(a["debrief_check"])
    footer_note(doc)
    return doc

def checklist(a):
    sp=STEPS[a["num"]]; t=TOPIC_BY_NUM[a["topic"]]
    doc=new_doc(); banner(doc); head(doc,a,"CHECKLIST"); meta_chips(doc,a); name_date(doc)

    doc.add_heading(f"Self-check — {a['title']}",level=2)
    doc.add_paragraph("Tick each item you have completed. Any blank is a gap to close before "
                      "your group presents at the debrief.")
    tbl=doc.add_table(rows=1,cols=4); tbl.style="Table Grid"
    h=tbl.rows[0].cells
    for i,x in enumerate(["#","Criterion","Evidences","✓"]):
        h[i].text=""; r=h[i].paragraphs[0].add_run(x)
        r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(h[i],BRAND)
    for i,item in enumerate(sp["checklist"],1):
        c=tbl.add_row().cells
        c[0].text=""; c[0].paragraphs[0].add_run(str(i)).font.size=Pt(9.5)
        c[1].text=""; c[1].paragraphs[0].add_run(item).font.size=Pt(9.5)
        c[2].text=""; c[2].paragraphs[0].add_run(t["code"]).font.size=Pt(9.5)
        c[3].text=""
        if i%2==0:
            for cc in c: shade(cc,LIGHT)
    doc.add_paragraph()

    doc.add_heading("Debrief check",level=2)
    p=doc.add_paragraph(); r=p.add_run(a["debrief_check"]); r.font.size=Pt(10.5)

    doc.add_heading("Feedback",level=2)
    for q in ["What did you find hardest in this activity?",
              "What is the ONE thing you will do differently at work?"]:
        p=doc.add_paragraph(); p.space_after=Pt(2); p.space_before=Pt(8)
        r=p.add_run(q); r.bold=True; r.font.size=Pt(10)
        blanks(doc,2)
    footer_note(doc)
    return doc

# ------------------------------------------------------------------ pdf
_SOFFICE=shutil.which("soffice") or shutil.which("libreoffice")

def to_pdf(docx_path):
    """Convert a DOCX to PDF beside it. Returns True on success."""
    if not _SOFFICE: return False
    outdir=os.path.dirname(docx_path)
    with tempfile.TemporaryDirectory() as tmp:
        profile=os.path.join(tmp,"lo")
        try:
            subprocess.run([_SOFFICE,"--headless",
                            f"-env:UserInstallation=file://{profile}",
                            "--convert-to","pdf","--outdir",outdir,docx_path],
                           check=True,capture_output=True,timeout=240)
        except Exception:
            return False
    return os.path.exists(os.path.splitext(docx_path)[0]+".pdf")

# ------------------------------------------------------------------ indexes
def readme_md():
    by_lu={}
    for a in ACT: by_lu.setdefault(a["topic"],[]).append(a)
    L=[f"# {C.TITLE} — Activities",""]
    L.append("17 progressive in-class activities across 4 Learning Units. Every activity is a "
             "**case study or role play** built on one continuous scenario — Nimbus Wellness Pte Ltd, "
             "a 25-person Singapore skincare brand running a 5th-anniversary relaunch — and builds "
             "toward the open-book Case Study assessment on Day 2.")
    L.append("")
    L.append("## What each activity folder contains")
    L.append("")
    L.append("| Artefact | What it is | Who uses it |")
    L.append("|---|---|---|")
    L.append("| `ANN-Facilitator-Guide-*.docx` / `.pdf` | Purpose, materials, set-up, scenario, the "
             "numbered run-sheet with facilitator notes, timing, debrief and success criterion | Trainer |")
    L.append("| `ANN-Learner-Worksheet-*.docx` / `.pdf` | The scenario, the steps, and a working space "
             "to capture the artefact and prompt answers | Learner |")
    L.append("| `ANN-Checklist-*.docx` / `.pdf` | Criteria table to tick before the debrief, plus feedback | Learner group |")
    L.append("")
    L.append("> The **detailed step-by-step instructions live here and in the Learner Guide — never on "
             "the slides.** The trainer facilitates the scenario, roles and discussion from the deck "
             "while learners work from these sheets.")
    L.append("")
    L.append("## Course structure")
    L.append("")
    for t in C.TOPICS:
        acts=by_lu[t["num"]]
        L.append(f"- **{t['code']} — {t['title']}** (Activities {acts[0]['num']}–{acts[-1]['num']}): {t['subtitle']}")
    L.append("")
    L.append("## Activities")
    L.append("")
    L.append("| # | LU | Activity | Type | Duration | Folder |")
    L.append("|---|---|---|---|---|---|")
    for a in ACT:
        t=TOPIC_BY_NUM[a["topic"]]; d=folder_for(a)
        L.append(f"| {a['num']} | {t['code']} | {a['title']} | {TYPE_LABEL.get(a.get('case_type'),'Workshop')} | "
                 f"{a['duration']} | `{d}/` |")
    L.append("")
    L.append("See [tools.md](tools.md) for the frameworks used across the activities.")
    L.append("")
    L.append(f"*{C.TITLE} · {C.COURSE_CODE} · Version {C.VERSION} · © 2026 {C.ORG}*")
    L.append("")
    return "\n".join(L)

def tools_md():
    L=["# Customer-Centric Branding Toolkit","",f"*{C.TITLE} · {C.COURSE_CODE}*",""]
    L.append("## Frameworks used in the activities")
    L.append("")
    L.append("| Framework | What it does | Used in |")
    L.append("|---|---|---|")
    for r in [
        ("Mendelow's Power-Interest Grid","Plots stakeholders by power and interest, then matches an engagement approach to each quadrant.","Activity 1"),
        ("Audience segmentation & messaging map","Profiles external audience segments and tailors message + channel per segment.","Activity 2"),
        ("The 5 C's of communication","Clear, Concise, Concrete, Correct, Courteous — a message quality test.","Activities 2, 9"),
        ("Hard/soft brand attribute mapping","Separates product facts (hard) from the feelings they create (soft).","Activity 3"),
        ("Keller's CBBE Pyramid","Locates a brand across Salience, Performance/Imagery, Judgments/Feelings, Resonance.","Activity 3"),
        ("Digital reputation audit","Samples mentions/reviews across platforms and codes sentiment into recurring themes.","Activities 4, 13"),
        ("Perception-gap analysis","Compares internal belief against customer feedback to find where perception diverges.","Activity 5"),
        ("Technical vs Value Quality","Separates what the product does from how it feels to deal with the brand.","Activity 5"),
        ("The Advocacy Loop","Traces satisfied → loyal → advocate → new customer, and finds where the loop breaks.","Activity 6"),
        ("PESO model","Sorts campaign activity into Paid, Earned, Shared and Owned media.","Activities 7, 12"),
        ("3 P's active listening","Presence, Patience, Paraphrasing — hearing the emotion behind the words.","Activity 8"),
        ("Brand guidelines framework","Defines values, visual identity and voice from an audit of existing touchpoints.","Activity 9"),
        ("Brand activation planning canvas","Plans a campaign — audience, message, channels, experiential element, recall metric.","Activity 11"),
        ("Leading vs lagging indicators","Separates early-warning metrics from after-the-fact confirmation metrics.","Activity 13"),
        ("Brand health scorecard","Rates Awareness, Loyalty, Financial Impact and Trust against evidence.","Activity 14"),
        ("Crisis holding statement","Acknowledge / act / update — speaking before root cause is confirmed.","Activity 15"),
        ("SMART KPI framework","Defines measure, target, owner, data source and review cadence.","Activity 16"),
        ("AMEC measurement framework","Inputs → Outputs → Outcomes → Impact for campaign evaluation.","Activity 17"),
    ]:
        L.append(f"| {r[0]} | {r[1]} | {r[2]} |")
    L.append("")
    L.append("## SMART KPI quick reference")
    L.append("")
    L.append("| Letter | Meaning |")
    L.append("|---|---|")
    for l,m in [("S","Specific — states exactly what is being measured"),
                ("M","Measurable — has a number or verifiable state"),
                ("A","Achievable — realistic given resources and timeframe"),
                ("R","Relevant — tied to the brand's actual goals"),
                ("T","Time-bound — has a deadline or review cadence")]:
        L.append(f"| **{l}** | {m} |")
    L.append("")
    L.append("## The 5 C's of customer communication")
    L.append("")
    L.append("| C | Test |")
    L.append("|---|---|")
    for l,m in [("Clear","Could the reader act on this without asking a question?"),
                ("Concise","Is every word doing work?"),
                ("Concrete","Are there specifics, not generalities?"),
                ("Correct","Is it accurate and error-free?"),
                ("Courteous","Would you be happy to receive it?")]:
        L.append(f"| **{l}** | {m} |")
    L.append("")
    L.append(f"*{C.TITLE} · {C.COURSE_CODE} · © 2026 {C.ORG}*")
    L.append("")
    return "\n".join(L)

# ------------------------------------------------------------------ main
if __name__=="__main__":
    os.makedirs(ACTIVITIES,exist_ok=True)

    # clear the previous flat *.md sheets — the folder layout is now authoritative
    for old in sorted(os.listdir(ACTIVITIES)):
        if re.match(r"^activity-\d{2}-.*\.md$",old):
            os.remove(os.path.join(ACTIVITIES,old)); print("Removed flat sheet",old)

    for name,text in (("README.md",readme_md()),("tools.md",tools_md())):
        with open(os.path.join(ACTIVITIES,name),"w",encoding="utf-8") as f: f.write(text)
        print("Wrote",name)

    pdf_ok=pdf_skip=0
    for a in ACT:
        d=os.path.join(ACTIVITIES,folder_for(a)); os.makedirs(d,exist_ok=True)
        ts=title_slug(a["title"]); n=f"A{a['num']:02d}"
        for kind,builder in (("Facilitator-Guide",facilitator_guide),
                             ("Learner-Worksheet",learner_worksheet),
                             ("Checklist",checklist)):
            path=os.path.join(d,f"{n}-{kind}-{ts}.docx")
            builder(a).save(path)
            if to_pdf(path): pdf_ok+=1
            else: pdf_skip+=1
        print(f"Wrote {folder_for(a)}/  (3 sheets)")

    print(f"\nPDFs rendered: {pdf_ok}   failed/skipped: {pdf_skip}")
    if pdf_skip and not _SOFFICE:
        print("NOTE: soffice/libreoffice not found — PDFs were not rendered.")
