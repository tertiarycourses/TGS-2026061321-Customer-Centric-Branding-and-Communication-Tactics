#!/usr/bin/env python3
"""Generate the WA (Written Assessment) + CS (Case Study) instruments for
Customer-Centric Branding and Communication Tactics, as four DOCX files:
question papers + model-answer/marking guides. WSQ house cover page (via
prodoc, no version-control record) — no multiple choice, every item
open-ended. WA tests knowledge from the slides/concepts; CS is one coherent
scenario whose model answers are drawn from the in-class activities' build
deliverable and discussion prompts.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4
import prodoc
REPO=os.path.dirname(os.path.dirname(HERE)); ASSETS=os.path.join(REPO,"courseware","assets")
OUT=os.path.join(REPO,"assessment")

Q_VER = A_VER = "v3"

DARK=RGBColor(0x16,0x1B,0x26); BRAND=RGBColor(0x1F,0x6F,0xEB); GREY=RGBColor(0x55,0x5B,0x66)

def new_doc():
    d=Document(); n=d.styles["Normal"]; n.font.name="Arial"; n.font.size=Pt(11)
    prodoc.style_headings(d); return d

def line(d,text="",bold=False,size=11,color=DARK,after=6,align=None):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    if align is not None: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.color.rgb=color; r.font.name="Arial"
    return p

def runs(d,segments,after=6,style=None):
    p=d.add_paragraph(style=style); p.paragraph_format.space_after=Pt(after)
    for text,bold in segments:
        r=p.add_run(text); r.bold=bold; r.font.size=Pt(11); r.font.name="Arial"; r.font.color.rgb=DARK
    return p

def bullet(d,text): return runs(d,[(text,False)],after=3,style="List Bullet")
def numbered(d,text): return runs(d,[(text,False)],after=3,style="List Number")

def _shade(cell,hexfill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr=cell._tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hexfill)
    tcPr.append(shd)

def _hyperlink(paragraph,url,text):
    """A real clickable hyperlink (python-docx has no API for this)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    part=paragraph.part
    r_id=part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hl=OxmlElement("w:hyperlink"); hl.set(qn("r:id"),r_id)
    new_run=OxmlElement("w:r"); rPr=OxmlElement("w:rPr")
    col=OxmlElement("w:color"); col.set(qn("w:val"),"1F6FEB"); rPr.append(col)
    u=OxmlElement("w:u"); u.set(qn("w:val"),"single"); rPr.append(u)
    new_run.append(rPr)
    t=OxmlElement("w:t"); t.text=text; new_run.append(t)
    hl.append(new_run); paragraph._p.append(hl)

def candidate_block(d,minutes,instructions):
    """Page 2 of every question paper: Trainee Information, Instructions to
    Candidate (with the clickable LMS link) and the Grading block — and nothing
    else. The scenario/questions then start on page 3."""
    line(d,"A: Trainee Information",bold=True,size=12,after=4)
    t=d.add_table(rows=3,cols=2); t.style="Table Grid"
    for i,lbl in enumerate(("Trainee Name (as per NRIC)",
                            "Last 3 digits and alphabet of NRIC / FIN","Date")):
        c=t.rows[i].cells[0]; c.text=""
        r=c.paragraphs[0].add_run(lbl); r.bold=True; r.font.size=Pt(10)
        _shade(c,"F5F8FC")
        t.rows[i].cells[1].text=""
    line(d,"",after=8)

    line(d,"B: Instructions to Candidate",bold=True,size=12,after=4)
    for t_ in instructions: line(d,t_,after=3)
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(3)
    r=p.add_run("Submission: upload your completed answers to the LMS at ")
    r.font.size=Pt(11); r.font.name="Arial"
    _hyperlink(p,"https://lms-tms.tertiaryinfotech.com/","https://lms-tms.tertiaryinfotech.com/")
    line(d,"",after=8)

    line(d,"C: Grading",bold=True,size=12,after=4)
    g=d.add_table(rows=3,cols=4); g.style="Table Grid"
    cells=[("Grade","(C / NYC)"),("Assessor Name",""),
           ("Assessor NRIC",""),("Date",""),("Signature",""),("",""),]
    labels=[("Grade (C / NYC)",""),("Assessor Name",""),("Assessor NRIC",""),
            ("Date",""),("Signature",""),("","")]
    for i,(lbl,_v) in enumerate(labels):
        row=g.rows[i//2]; c=row.cells[(i%2)*2]
        c.text=""
        if lbl:
            r=c.paragraphs[0].add_run(lbl); r.bold=True; r.font.size=Pt(10)
            _shade(c,"F5F8FC")
        row.cells[(i%2)*2+1].text=""
    d.add_page_break()

def answer_lines(d,n=3):
    """A bordered answer box sized by the number of lines expected."""
    line(d,"Answer:",bold=True,color=GREY,after=2)
    t=d.add_table(rows=1,cols=1); t.style="Table Grid"
    cell=t.rows[0].cells[0]; cell.text=""
    for _ in range(max(1,n)):
        p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(10)
    line(d,"",after=8)

def cover(d,kind):
    prodoc.add_cover_page(d,kind,C.TITLE,Q_VER.lstrip("v"),
                          org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                          course_logo=None,course_code=C.COURSE_CODE)

# ================================================================ WA — content
# One open-ended question per concept taught (K1..K16). (question, answer_lines, source, [model points])
_Q_STEMS=[
    "Distinguish between internal and external stakeholders of a brand, and explain why their level of "
    "influence over the brand can differ.",
    "Name the different types of external audience a brand must communicate with, and explain why each "
    "needs its own communication approach and channel.",
    "When drafting a branding design or idea, why is it important to highlight the product or service's "
    "attributes and benefits rather than just its features?",
    "Describe how you would assess an organisation's reputation on social media and other platforms, and "
    "explain why this is the starting point for a branding decision.",
    "Explain what is meant by a \"perception gap\" between how a company sees its brand and how customers "
    "actually experience it, and why this gap matters.",
    "Explain how customers can influence a brand's reputation, and why their voice can have more impact "
    "than a paid campaign.",
    "Why is it important to document customer reception and the outcome of a branding campaign, rather "
    "than relying on general impressions?",
    "Describe the three elements of active listening (Presence, Patience, Paraphrasing), and explain how "
    "they help reveal a customer's true perspective of an organisation.",
    "List the basic elements that make up a brand's identity, and explain why consistency across these "
    "elements matters.",
    "Explain the role that branding plays within the marketing mix, and how it differs from a single "
    "marketing campaign.",
    "Describe how branding campaigns, events and activities can be used to increase brand awareness.",
    "Explain why a public-relations campaign must be aligned to the brand positioning strategy, "
    "operational plan and budget.",
    "Identify the different types of metrics used to measure an organisation's reputation, and explain "
    "why they must be tracked across multiple platforms.",
    "List indicators that show a branding effort has been successful, beyond just awareness.",
    "Describe the public-relations tactics a brand can use to extend its message beyond paid channels.",
    "Explain how monitoring a brand against SMART KPIs turns branding into a continuous-improvement "
    "process.",
]
# Each question carries its own K code, printed on the paper AND repeated in the
# key, so every K1..K16 is visibly covered by the WA.
WRITTEN=[]
_qi=0
for t in C.TOPICS:
    for c in t["concepts"]:
        WRITTEN.append((_Q_STEMS[_qi], 3, f"{t['code']} — {t['title']}", [c], f"K{_qi+1}"))
        _qi+=1

def check_coverage():
    """Fail the build if any K or A is orphaned."""
    k_expected={f"K{i}" for i in range(1,len(WRITTEN)+1)}
    k_used={code for *_,code in WRITTEN}
    a_expected={f"A{i}" for i in range(1,len(CS_Q)+1)}
    a_used={code for _,code,_ in CS_Q}
    problems=[]
    if k_expected-k_used: problems.append(f"WA missing {sorted(k_expected-k_used)}")
    if a_expected-a_used: problems.append(f"CS missing {sorted(a_expected-a_used)}")
    print("  K coverage:", ", ".join(f"{c}->Q{i}" for i,(*_,c) in enumerate(WRITTEN,1)))
    print("  A coverage:", ", ".join(f"{c}->{lbl}" for lbl,c,_ in CS_Q))
    if problems:
        raise SystemExit("COVERAGE FAILURE: "+"; ".join(problems))
    return True

def build_written_paper():
    d=new_doc(); cover(d,"WRITTEN ASSESSMENT (SAQ)")
    candidate_block(d,60,[
        "1. This is an individual, open-book assessment of underpinning knowledge.",
        "2. A total of 1 hour is given to complete this Written Assessment.",
        f"3. There are {len(WRITTEN)} questions. Answer ALL questions in your own words in the space provided.",
        "4. All questions are based on the concepts taught in LU1–LU4.",
        "5. You need to answer all questions correctly to be assessed as Competent.",
    ])
    line(d,"Written Assessment (SAQ) — Knowledge",bold=True,size=15,color=BRAND,after=2,align=AL.CENTER)
    line(d,f"Course Code: {C.COURSE_CODE}",size=10,color=GREY,after=12,align=AL.CENTER)
    line(d,"D: Answer all questions.",bold=True,size=12,after=6)
    for i,(q,nlines,_,_,code) in enumerate(WRITTEN,1):
        runs(d,[(f"Q{i}. ",True),(q,False),(f"  ({code})",True)],after=2)
        answer_lines(d,nlines)
    out=os.path.join(OUT,f"WA (SAQ) - {C.TITLE} - {Q_VER}.docx"); d.save(out); return out

def build_written_answers():
    d=new_doc(); cover(d,"ANSWER KEY — WRITTEN ASSESSMENT (SAQ)")
    line(d,"Written Assessment (SAQ) — Answer Key & Marking Guide",bold=True,size=15,color=BRAND,after=2,align=AL.CENTER)
    line(d,f"Course Code: {C.COURSE_CODE}",size=10,color=GREY,after=12,align=AL.CENTER)
    line(d,"Mark each answer against the model points below — award the mark where the candidate covers "
           "the key ideas; wording will vary. All items are taught in LU1–LU4 (source shown in brackets).",after=8)
    for i,(q,_,src,points,code) in enumerate(WRITTEN,1):
        runs(d,[(f"Q{i}. ",True),(q,False),(f"  ({code})",True),(f"   [{src}]",False)],after=2)
        line(d,"Suggestive answers (not exhaustive):",bold=True,color=GREY,after=2)
        for p in points: bullet(d,p)
        line(d,"",after=4)
    out=os.path.join(OUT,f"Answer to WA (SAQ) - {C.TITLE} - {Q_VER}.docx"); d.save(out); return out

# ================================================================ CASE STUDY — content
SCENARIO=(
    "Nimbus Wellness Pte Ltd is a Singapore-based boutique skincare and wellness brand with 25 staff "
    "across Retail, Marketing, Customer Service and Operations. Since a competitor's viral campaign "
    "six months ago, Nimbus has seen its social-media engagement fall and a rise in mixed reviews "
    "questioning whether the brand still understands its customers. Leadership wants a customer-centric "
    "branding and communication plan built entirely using the techniques practised in class, before the "
    "brand's 5th-anniversary relaunch campaign in three months."
)
SCENARIO_POINTS=[
    "Marketing believes the product quality is unchanged, but online reviews increasingly say the brand "
    "\"doesn't feel personal anymore\" — a possible perception gap.",
    "The founder wants a refreshed brand identity for the relaunch, but different teams describe the "
    "brand's values inconsistently in customer-facing material.",
    "The relaunch needs a multi-channel awareness campaign with a bounded PR budget approved by the CEO.",
    "The Board wants ongoing proof, after the relaunch, that the campaign actually improved the brand's "
    "reputation and not just its impressions.",
]
SCENARIO_TAIL=(
    "Using exactly the techniques you practised in the in-class activities, complete the four tasks "
    "below for Nimbus Wellness."
)

CS_Q=[
    ("Question 1","A1",
     "Stakeholders and brand design. Identify Nimbus's internal and external stakeholders and audiences, "
     "and draft a brand-attribute concept (hard + soft attributes) for the relaunch that reflects the "
     "product's real benefits. State how you would assess Nimbus's current reputation on social media."),
    ("Question 2","A2",
     "Customer influence. Using an active-listening approach, explain how you would investigate the "
     "\"doesn't feel personal anymore\" perception gap, and how you would document customer reception once "
     "you have gathered the feedback."),
    ("Question 3","A3",
     "Branding in marketing. Propose a brand-activation plan for the relaunch campaign — channels, core "
     "message, and one experiential element — and show how the PR budget would be allocated across "
     "earned, owned and paid channels in line with the brand positioning."),
    ("Question 4","A4",
     "Branding effectiveness. Design a KPI monitoring framework the Board can use after the relaunch, and "
     "explain what PR tactics and crisis-response steps Nimbus should have ready if the campaign attracts "
     "negative coverage."),
]

def _acts(nums): return [a for a in ACT if a["num"] in nums]

CS_A=[
    ("Question 1","A1","Activities 1–4 — Stakeholder Influence Mapping, Audience Mapping Workshop, "
     "Brand Attribute Mapping, Digital Reputation Audit Exercise",
     [s for a in _acts([1,2,3,4]) for s in ([f"{a['title']}: {a['build']}."]+a['discussion_prompts'])]),
    ("Question 2","A2","Activities 5–8 — Brand Perception Audit, Brand Advocacy Assessment, "
     "Campaign Documentation Audit, Customer Perspective Analysis",
     [s for a in _acts([5,6,7,8]) for s in ([f"{a['title']}: {a['build']}."]+a['discussion_prompts'])]),
    ("Question 3","A3","Activities 9–12 — Brand Audit Workshop, Brand-Marketing Audit Exercise, "
     "Brand Activation Planning Workshop, PR Campaign Budget Planning",
     [s for a in _acts([9,10,11,12]) for s in ([f"{a['title']}: {a['build']}."]+a['discussion_prompts'])]),
    ("Question 4","A4","Activities 13–17 — Platform Reputation Audit, Brand Health Assessment, "
     "PR Crisis Response Plan, KPI Dashboard Design, PR Campaign Audit Exercise",
     [s for a in _acts([13,14,15,16,17]) for s in ([f"{a['title']}: {a['build']}."]+a['discussion_prompts'])]),
]

def case_study_block(d):
    line(d,"D: Case Study",bold=True,size=12,after=4)
    line(d,"Scenario:",bold=True,after=2)
    line(d,SCENARIO,after=4)
    for pt in SCENARIO_POINTS: bullet(d,pt)
    line(d,SCENARIO_TAIL,after=8)

def build_case_paper():
    d=new_doc(); cover(d,"CASE STUDY (CS)")
    candidate_block(d,60,[
        "1. This is an individual, open-book practical assessment.",
        "2. A total of 1 hour is given to complete this Case Study assessment.",
        f"3. There are {len(CS_Q)} tasks. Answer ALL tasks, based on the Nimbus Wellness case study.",
        "4. Write your answer as you would present it to Nimbus's leadership team.",
        "5. All tasks reuse the techniques you practised in the in-class activities.",
    ])
    line(d,"Case Study (CS) — Practical",bold=True,size=15,color=BRAND,after=2,align=AL.CENTER)
    line(d,f"Course Code: {C.COURSE_CODE}",size=10,color=GREY,after=12,align=AL.CENTER)
    case_study_block(d)
    for label,crit,body in CS_Q:
        runs(d,[(f"{label} ({crit}): ",True),(body,False)],after=4)
        answer_lines(d,6)
    out=os.path.join(OUT,f"CS Assessment - {C.TITLE} - {Q_VER}.docx"); d.save(out); return out

def build_case_answers():
    d=new_doc(); cover(d,"ANSWER KEY — CASE STUDY (CS) ASSESSMENT")
    line(d,"Case Study (CS) Assessment — Model Answers (Marking Guide)",bold=True,size=15,color=BRAND,after=2,align=AL.CENTER)
    line(d,f"Course Code: {C.COURSE_CODE}",size=10,color=GREY,after=12,align=AL.CENTER)
    line(d,"Note to assessor:",bold=True,after=2)
    line(d,"Each task maps to one ability criterion (A1–A4) and to the in-class activities the candidate "
           "completed — the model answer is drawn from the discussion points those activities worked through. "
           "Award Competent (C) where the candidate produces (or correctly describes) a plan that covers the "
           "points below; exact wording will vary.",after=8)
    case_study_block(d)
    for label,crit,act_ref,points in CS_A:
        line(d,f"{label} ({crit}) — Model Answer",bold=True,size=12,color=BRAND,after=2)
        runs(d,[("Reference: ",True),(act_ref,False)],after=4)
        line(d,"Suggestive answers (not exhaustive):",bold=True,color=GREY,after=2)
        for p in points: bullet(d,p)
        line(d,"",after=4)
    out=os.path.join(OUT,f"Answer to CS Assessment - {C.TITLE} - {A_VER}.docx"); d.save(out); return out

if __name__=="__main__":
    os.makedirs(OUT,exist_ok=True)
    print("K/A coverage check:")
    check_coverage()
    for fn in (build_written_paper,build_written_answers,build_case_paper,build_case_answers):
        print("Wrote",fn())
