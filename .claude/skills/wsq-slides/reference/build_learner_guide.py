#!/usr/bin/env python3
"""Generate the AZ-104 Learner Guide as BOTH a Markdown mirror (LG-*.md at repo
root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll build · Step-by-step
with commands · Test it), plus setup, exam-prep and glossary. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned with
the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5
import prodoc
REPO=os.path.dirname(os.path.dirname(HERE)); ASSETS=os.path.join(REPO,"courseware","assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  "It provides step-by-step instructions for all 26 hands-on labs, organised by the five official "
  "AZ-104 exam skill areas. Every lab maps to a published exam objective and is completed in the "
  "Azure Portal together with Azure Cloud Shell (Azure CLI and PowerShell).")
p("Use this guide alongside the course slides and the lab files in the labs/ folder of the course "
  "repository. Each lab creates its own resource group (rg-az104-labNN); always run the clean-up "
  "step at the end of a lab to avoid unnecessary Azure charges.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Before You Start — Environment Setup")
h3("What you need")
bullets([
 "An Azure subscription — an Azure free account or an instructor-provided subscription.",
 "Contributor access on the subscription; Global Administrator / User Administrator on the Microsoft Entra tenant for the Topic 1 identity labs.",
 "A modern browser for the Azure Portal (https://portal.azure.com) and Cloud Shell (https://shell.azure.com).",
])
h3("Launch Azure Cloud Shell")
p("Cloud Shell is a browser terminal with the az CLI, Az PowerShell, azcopy and Bicep pre-installed — nothing to install locally. Open it from the >_ icon in the portal, accept the storage prompt on first launch, then choose Bash (for az) or PowerShell (for Az).")
code("az version            # Azure CLI\naz account show        # confirm your subscription\naz account set --subscription \"<name-or-id>\"   # if you have more than one")
h3("Conventions used in every lab")
bullets([
 "Each lab uses its own resource group named rg-az104-labNN.",
 "Pick one region (e.g. eastus) and use it consistently across a lab.",
 "Globally-unique names (storage accounts, ACR, web apps) append a random suffix — change it if a name is taken.",
 "Run the Clean up step at the end of each lab: az group delete --name rg-az104-labNN --yes --no-wait",
])

# ---------------- per-topic, per-lab ----------------
TOPICS_BY_NUM={t["num"]:t for t in C.TOPICS}
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}  ({t['weighting']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets(t["concepts"])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Lab {a['num']} — {a['title']}")
        p(f"Exam objective: {a['objective']}.")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Azure services: {a['services']}.)")
        h3("Step-by-step")
        st=[]
        for i,(instr,cmd) in enumerate(a["steps"],1):
            st.append((instr,cmd))
        steps(st)
        h3("Test it")
        p(a["test"])
        note(f"Full commands (Portal + CLI + PowerShell) are in labs/lab-{a['num']:02d}-*.md. "
             f"Clean up when done: az group delete --name rg-az104-lab{a['num']:02d} --yes --no-wait")
        rule()

h1("Exam Preparation")
bullets([
 "First pass: do every lab via the Azure Portal, reading the References in each lab file.",
 "Second pass: redo the labs using only the Azure CLI until the command verbs are automatic.",
 "Review the 'Test it' check and the 'What you learned' bullets for any topic you find hard.",
 "Take the free Microsoft practice assessment for AZ-104.",
 "Passing score is 700/1000. Book the exam from your Microsoft Learn profile.",
])

h1("Glossary")
gl=[
 ("Resource group","A container that holds related Azure resources sharing a lifecycle."),
 ("Microsoft Entra ID","Azure's cloud identity service (authentication) — formerly Azure AD."),
 ("RBAC","Role-Based Access Control — grants a principal a role at a scope (authorization)."),
 ("Azure Policy","Rules that audit or enforce resource configuration for governance."),
 ("Storage account","A globally-unique namespace holding Blobs, Files, Queues and Tables."),
 ("SAS token","Shared Access Signature — a scoped, time-limited URL for storage access."),
 ("ARM / Bicep","Azure Resource Manager templates / the Bicep language for infrastructure as code."),
 ("VNet","Virtual Network — your private, isolated network in Azure, divided into subnets."),
 ("NSG","Network Security Group — stateful allow/deny rules filtering subnet or NIC traffic."),
 ("Azure Bastion","Managed service for browser RDP/SSH to VMs without a public IP."),
 ("Log Analytics / KQL","Azure Monitor's log store and its Kusto Query Language."),
 ("Recovery Services vault","The container Azure Backup and Site Recovery use to store recovery points."),
]
B.append(("dl",gl))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[(C.VERSION.lstrip("v"),C.VERSION_DATE,"Initial release — AZ-104 Learner Guide covering all 26 labs.",C.TRAINER)])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph(style="List Number"); para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
