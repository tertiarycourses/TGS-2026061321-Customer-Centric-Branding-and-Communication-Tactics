#!/usr/bin/env python3
"""Generate the AZ-104 Lesson Plan (LP) DOCX in the Tertiary house format.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 3-day schedule tables (9:00am-6:00pm, 8 training hours/day, 1h
lunch, tea within, final assessment Day 3 4:00pm). Topics/labs come from
course_data + the domain data files so the LP stays aligned with the deck,
guide and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5
import prodoc
REPO=os.path.dirname(os.path.dirname(HERE)); ASSETS=os.path.join(REPO,"courseware","assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

def lab_titles(nums):
    return "; ".join(f"Lab {a['num']}: {a['title']}" for a in ACT if a['num'] in nums)

# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, activity_text)  kind: admin/topic/lab/break/lunch/assess/recap
SCHEDULE = {
 1: ("Manage Azure Identities, Governance & Storage", [
    ("9:00","9:30",30,"admin","Welcome, course introduction, ground rules and mandatory digital attendance (AM)"),
    ("9:30","10:30",60,"topic","Topic 1 — Manage Azure Identities and Governance: Entra ID, RBAC, Azure Policy, subscriptions, resource groups, locks, tags and cost (concepts + demo)"),
    ("10:30","10:45",15,"break","Tea break"),
    ("10:45","13:00",135,"lab","Hands-on: "+lab_titles([1,2,3,4])),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:30",90,"lab","Hands-on: "+lab_titles([5,6])+". Topic 2 — Implement and Manage Storage: storage accounts, redundancy, encryption (concepts)"),
    ("15:30","15:45",15,"break","Tea break"),
    ("15:45","17:45",120,"lab","Hands-on: "+lab_titles([7,8,9,10,11])),
    ("17:45","18:00",15,"recap","Day 1 recap, Q&A and PM digital attendance"),
 ]),
 2: ("Deploy Compute Resources & Virtual Networking", [
    ("9:00","9:15",15,"recap","Day 1 recap and mandatory digital attendance (AM)"),
    ("9:15","10:30",75,"topic","Topic 3 — Deploy and Manage Azure Compute Resources: ARM/Bicep, VMs, scale sets, containers, App Service (concepts + demo)"),
    ("10:30","10:45",15,"break","Tea break"),
    ("10:45","13:00",135,"lab","Hands-on: "+lab_titles([12,13,14,15])),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:30",90,"lab","Hands-on: "+lab_titles([16,17])),
    ("15:30","15:45",15,"break","Tea break"),
    ("15:45","17:45",120,"lab","Topic 4 — Implement and Manage Virtual Networking (concepts). Hands-on: "+lab_titles([18,19])),
    ("17:45","18:00",15,"recap","Day 2 recap, Q&A and PM digital attendance"),
 ]),
 3: ("Networking, Monitoring, Maintenance & Assessment", [
    ("9:00","9:15",15,"recap","Day 2 recap and mandatory digital attendance (AM)"),
    ("9:15","10:45",90,"lab","Hands-on: "+lab_titles([20,21,22])),
    ("10:45","11:00",15,"break","Tea break"),
    ("11:00","13:00",120,"lab","Topic 5 — Monitor and Maintain Azure Resources (concepts). Hands-on: "+lab_titles([23,24])),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:30",90,"lab","Hands-on: "+lab_titles([25,26])),
    ("15:30","15:45",15,"break","Tea break"),
    ("15:45","16:00",15,"assess","Briefing for Assessment"),
    ("16:00","17:00",60,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, open book"),
    ("17:00","18:00",60,"assess","Practical Performance (PP) — hands-on Azure tasks, 1 hour, open book. PM digital attendance"),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[(C.VERSION.lstrip("v"),C.VERSION_DATE,"Initial release — AZ-104 3-day lesson plan aligned to the 26 labs.",C.TRAINER)])
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration","3 days · 8 training hours per day (24 hours)"),
      ("Daily Timing","9:00 am – 6:00 pm (1-hour lunch; tea breaks within training time)"),
      ("Mode","Instructor-led, hands-on labs in the Azure Portal and Azure Cloud Shell"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "Final assessment is conducted on Day 3 from 4:00 pm.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","lab":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        if kind!="lunch": training+=mins
    # widths
    for row in tbl.rows:
        row.cells[0].width=Inches(1.15); row.cells[1].width=Inches(0.9); row.cells[2].width=Inches(4.75)
    p=doc.add_paragraph(); r=p.add_run(f"Total training time: {training} minutes ({training//60} hours)."); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    assert training==480, f"Day {day} training minutes = {training}, expected 480"

H("Lab Reference (aligned to exam skill areas)",1)
tt=doc.add_table(rows=0,cols=3); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Topic / Exam skill area","Weighting","Labs"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in ACT if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Topic {tp['code']}: {tp['title']}",bold=True,size=9.5,fill=TOPIC_FILL)
    set_cell(cells[1],tp["weighting"],size=9.5,fill=TOPIC_FILL)
    set_cell(cells[2],", ".join(f"Lab {a['num']}" for a in acts),size=9.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
