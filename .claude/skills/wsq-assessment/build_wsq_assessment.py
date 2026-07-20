#!/usr/bin/env python3
"""
WSQ ASSESSMENT GENERATOR (TEMPLATE) — produces two instruments as DOCX, each as a
question paper + a model-answer / marking guide:

  1. Written Assessment (WA)  — tests KNOWLEDGE. ALL questions are OPEN-ENDED
     short-answer (NO multiple choice). Every item is drawn from the concept
     modules / slides taught in class.
  2. Case Study (CS) Assessment — tests PRACTICAL ability. ONE coherent case
     study built from the in-class activities; the model answers ARE the exact
     build steps from the hands-on labs / practicals.

HOW TO USE: copy this file into the course repo (e.g. courseware/build_assessment.py),
then edit the CONFIG block and the four content lists — WRITTEN, SCENARIO*, CS_Q, CS_A.
Keep every Written item open-ended; never add MCQ options. Run: python3 <thisfile>.

The example content below is a worked sample — replace it with the target
course's own modules, slides and labs.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

# ============================== CONFIG — EDIT THIS ==============================
REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(REPO, "assessemnt")          # output folder for the four DOCX
TITLE = "Microsoft Copilot Studio & Power Automate for Business Workflow Automation"
COURSE_CODE = "TGS-2022017524"
# ===============================================================================
DARK = RGBColor(0x16, 0x1B, 0x26); BRAND = RGBColor(0x1F, 0x6F, 0xEB); GREY = RGBColor(0x55, 0x5B, 0x66)

# ================================================================ doc helpers
def new_doc():
    d = Document(); n = d.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11); return d

def line(d, text="", bold=False, size=11, color=DARK, after=6, align=None):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Arial"
    return p

def runs(d, segments, after=6, style=None):
    """segments: list of (text, bold)."""
    p = d.add_paragraph(style=style); p.paragraph_format.space_after = Pt(after)
    for text, bold in segments:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(11); r.font.name = "Arial"; r.font.color.rgb = DARK
    return p

def bullet(d, text):
    return runs(d, [(text, False)], after=3, style="List Bullet")

def numbered(d, text):
    return runs(d, [(text, False)], after=3, style="List Number")

def title_block(d, subtitle):
    line(d, TITLE, bold=True, size=15, color=BRAND, after=2, align=AL.CENTER)
    line(d, subtitle, bold=True, size=12, color=DARK, after=2, align=AL.CENTER)
    line(d, f"Course Code: {COURSE_CODE}", size=10, color=GREY, after=12, align=AL.CENTER)

def candidate_block(d, hours, instructions):
    line(d, "A: Candidate Information", bold=True, size=12, after=4)
    line(d, "Trainee Name (as per NRIC): _______________________________", after=4)
    line(d, "Last 3 digits and alphabet of NRIC / FIN: _________________", after=4)
    line(d, "Date: __________________", after=10)
    line(d, "B: Instructions to Candidate", bold=True, size=12, after=4)
    for t in instructions: line(d, t, after=3)
    line(d, "", after=4)

def official_use(d):
    line(d, "____________________________________________________________________________", color=GREY, after=6)
    line(d, "For Official Use Only", bold=True, after=4)
    line(d, "Grade: _____ (C / NYC)", after=4)
    line(d, "Assessor Name: _______________\t\tAssessor NRIC: _____________", after=4)
    line(d, "Date: ________________________\t\tSignature: _________________", after=4)

def answer_lines(d, n=3):
    line(d, "Answer:", bold=True, color=GREY, after=2)
    for _ in range(n):
        line(d, "_______________________________________________________________", color=GREY, after=4)

# ================================================================ WRITTEN — content
# All open-ended short-answer questions (no multiple choice).
# (question, answer_lines, source, [model answer points])
WRITTEN = [
    ("List the three traits that make a task a good candidate for automation.", 2, "Module 1",
     ["Repetitive — done the same way many times.",
      "Rule-based — clear “if this, then that” logic.",
      "Time-consuming — manual copying, chasing, or re-typing."]),
    ("State the three-part mental model of a workflow used throughout this course.", 2, "Module 1",
     ["Trigger → Actions → Output.",
      "An event starts it (trigger), the flow performs steps (actions), and data is produced/passed along (output)."]),
    ("List the four building blocks of workflow logic, with a one-word example of each.", 4, "Module 1",
     ["Trigger — what starts the workflow (e.g. an email arrives / a form is submitted).",
      "Actions — what the workflow does (e.g. send an email, add a row, start an approval).",
      "Outputs — the data the workflow produces or passes along (e.g. the new row, the approval outcome).",
      "Steps — the ordered sequence as a whole, including conditions (if/else) and loops."]),
    ("Name three flow types covered in the course and the event that starts each one.", 3, "Module 2",
     ["Instant / manual — a person presses Run (Labs 1–3).",
      "Scheduled — a Recurrence timetable is reached (Lab 4).",
      "Automated — an event such as a new form response, new email, or new file (Lab 5, Day 3).",
      "Agent flow — a Copilot Studio agent calls the flow as a tool (Day 2–3).  [any three]"]),
    ("Which trigger lets a Copilot Studio agent run a Power Automate flow as a tool? Name it exactly.", 2, "Module 2 / 3",
     ["“When an agent calls the flow.”",
      "The flow then ends with a “Respond to the agent” action to return a result."]),
    ("In Power Automate, Excel actions can only read or write data that is inside what structure? Why does this matter?", 2, "Module 2 / Lab 2",
     ["A named Table (Insert → Table) — not loose cells.",
      "Actions like “Add a row into a table” only see data inside a Table; this is a classic first-time gotcha."]),
    ("What does an “Unauthorized” error usually mean when a Send an email action runs, and how do you fix it?", 2, "Module 2",
     ["The signed-in account lacks rights for that action (e.g. it is not mailbox-enabled), or the connection is broken/expired.",
      "Fix: reconnect the connector with a valid, mailbox-enabled tenant account. Golden rule: green = ready, red = reconnect first."]),
    ("In an approval step (Start and wait for an approval), who must the approver be, and what happens if you ignore this?", 2, "Module 2 / 4",
     ["The approver must be a real user in your Microsoft 365 tenant (yourself is fine for testing).",
      "An external / personal email address never receives the request, so the flow appears to hang."]),
    ("What technique lets an agent’s Knowledge answer from your own documents? Spell out the acronym and say what it prevents.", 2, "Module 3 / Lab 7",
     ["RAG — Retrieval-Augmented Generation: it retrieves relevant passages from your files and generates an answer grounded in them.",
      "Grounding the agent in your sources (and turning General knowledge off) prevents hallucination / made-up answers."]),
    ("What is the most reliable way to get structured output from an agent so a flow receives clean inputs?", 2, "Module 3 / Lab 9",
     ["Capture each answer into its own named variable using Ask a question nodes (e.g. customerName, product, quantity).",
      "Set numeric fields to Number and re-ask on missing/invalid input. Structured agent output = clean flow inputs."]),
    ("List the five components of a Copilot Studio agent and what each is for.", 5, "Module 3",
     ["Instructions — plain-language directions that shape the agent’s behaviour.",
      "Knowledge — documents / sites the agent can answer from (RAG).",
      "Topics — conversation flows triggered by what the user says.",
      "Tools / Actions — things the agent can do, including Power Automate flows.",
      "Variables — where captured answers are stored to pass onward."]),
    ("With generative orchestration, what does the agent use to decide which tool to call? What follows for how you write a tool?", 2, "Lab 8",
     ["It selects the tool from the tool’s name and description (there are no fixed trigger phrases).",
      "So write a clear, intent-rich description; if two tools overlap, disable the one you are not using so the agent doesn’t pick the wrong one."]),
    ("Module 4 says good orchestration has “no dead ends.” List the five elements every orchestrated workflow should have.", 5, "Module 4",
     ["A clear trigger — one starting event.",
      "Sequenced actions — each step’s output feeds the next step’s input.",
      "Branching — conditions route the process (approved vs rejected, high vs low value).",
      "Notifications — the right people are told at the right moments.",
      "A definite end state — the record is updated to its final status (Logged / Approved / Rejected / Completed)."]),
    ("Name the simple six-step method taught for designing a workflow before you build it.", 4, "Module 4",
     ["Write the one-sentence summary: “When ___ happens, do ___, then ___, and finally ___.”",
      "Identify the trigger (the “when”); list the actions in order (the “do / then / finally”).",
      "Mark the decision points (where it branches); define each end state (final status + who is notified).",
      "Build it, test the happy path, then test every branch."]),
]

def build_written_paper():
    d = new_doc()
    title_block(d, "Written Assessment (WA) — Knowledge")
    candidate_block(d, 1, [
        "1. This is an individual, open-book assessment of underpinning knowledge.",
        "2. A total of 45 minutes is given to complete this Written Assessment.",
        "3. Answer all questions in your own words in the space provided.",
        "4. All questions are based on the concepts taught in Modules 1–4.",
    ])
    line(d, "Answer all questions.", bold=True, size=12, after=6)
    for i, (q, nlines, _, _) in enumerate(WRITTEN, 1):
        runs(d, [(f"Q{i}. ", True), (q, False)], after=2)
        answer_lines(d, nlines)
    official_use(d)
    out = os.path.join(OUT, f"Written Assessment(WA) - {TITLE}.docx")
    d.save(out); return out

def build_written_answers():
    d = new_doc()
    title_block(d, "Written Assessment (WA) — Answer Key & Marking Guide")
    line(d, "Mark each answer against the model points below — award the mark where the candidate covers "
            "the key terms; wording will vary. All items are taught in Modules 1–4 (source shown in brackets).", after=8)
    for i, (q, _, src, points) in enumerate(WRITTEN, 1):
        runs(d, [(f"Q{i}. ", True), (q, False), (f"   [{src}]", False)], after=2)
        for p in points: bullet(d, p)
        line(d, "", after=4)
    out = os.path.join(OUT, f"Answers to Written Assessment(WA) - {TITLE}.docx")
    d.save(out); return out

# ================================================================ CASE STUDY — content
SCENARIO = (
    "ACME Pte Ltd is a Singapore-based office-supplies distributor with about 60 staff across "
    "Sales, Finance, Procurement and Operations. Today most of its everyday processes still run on "
    "email and shared spreadsheets:"
)
SCENARIO_POINTS = [
    "Customers email or call in sales enquiries; a salesperson replies manually and re-types the details into a spreadsheet.",
    "Customers place orders by email; Operations re-keys each order into an order sheet and writes a confirmation email by hand.",
    "Staff raise purchase requests by emailing their manager; approvals are chased over email and the status is tracked manually.",
    "Suppliers send invoices as PDF attachments; Finance saves them to a folder, seeks approval by email, and records the outcome in a spreadsheet.",
]
SCENARIO_TAIL = (
    "The manual approach causes lost enquiries, inconsistent acknowledgement emails, missed approval "
    "thresholds, and no reliable audit trail. ACME has Microsoft 365 (Outlook, Excel on OneDrive, "
    "Microsoft Forms), Power Automate and Copilot Studio available in a managed environment "
    "(“ACME Sandbox”), with the Approvals and Office 365 Outlook connectors enabled. Using exactly the "
    "techniques you built in the hands-on labs, build (or describe the build of) the automations below."
)

# CS questions: (label, criterion, prompt)
CS_Q = [
    ("Question 1", "A1",
     "Sales-enquiry workflow (Power Automate). Build a flow that captures a customer enquiry, logs it to "
     "an Excel table, and acknowledges the customer. State the trigger, the actions in order, the "
     "connector each step uses, and the final end state."),
    ("Question 2", "A2",
     "Purchase-request approval (Power Automate). Build a flow that captures a staff purchase request, "
     "logs it to an Excel table, routes requests above a cost threshold to a manager for approval, and "
     "notifies the requester of the outcome. Include the condition and both approved / rejected paths."),
    ("Question 3", "A3",
     "Customer assistant (Copilot Studio). Build an agent that (a) answers only from ACME’s own "
     "documents (knowledge / RAG), (b) captures an order as structured data (product, quantity, "
     "delivery address), and (c) has a tool/action so it can do real work."),
    ("Question 4", "A4",
     "Order-processing, end-to-end (agent + flow). Connect the agent from Q3 to a Power Automate flow so "
     "an order is logged, a confirmation is generated, large orders raise a restock alert, and the "
     "customer is notified. Describe the agent→flow connection, the data mapping, every outcome, and "
     "how you test it."),
]

# CS model answers = the lab build steps. (label, criterion, lab_ref, [numbered steps])
CS_A = [
    ("Question 1", "A1", "Labs 1, 2, 5, 11 — pattern: Capture → Log → Notify (Module 4)", [
        "Trigger: start with “Manually trigger a flow” to build and test, then swap in the real trigger "
        "— Microsoft Forms “When a new response is submitted” (Lab 5). (Module 2 tip: build manual first, swap the trigger last.)",
        "Capture the fields from the form response (name, company, product interest, contact email).",
        "Add a row into a table — Excel Online (Business), into a named “Leads” Table; map the fields and "
        "set the date with the fx editor: formatDateTime(utcNow(),'yyyy-MM-dd HH:mm') (Lab 2). Excel only sees data inside a Table.",
        "Generate a personalised acknowledgement with an AI prompt (Lab 11).",
        "Send an email (V2) — Office 365 Outlook — the acknowledgement to the customer (Lab 1); use dynamic content for name/product.",
        "Notify the assigned salesperson by a second Send an email (V2) (or Post message in Teams).",
        "End state: the lead is logged in Excel (audit trail) and the customer + salesperson are notified. "
        "Confirm every connection is green ✓ before testing, and check the row + emails in a test run (run history all green).",
    ]),
    ("Question 2", "A2", "Labs 2, 3, 14 — pattern: Request → Approve → Notify (Module 4)", [
        "Trigger / capture: a Microsoft Form or an agent captures requester, item, cost, justification. "
        "Set the cost field type to Number so it can be compared (Lab 9 / Lab 14).",
        "Add a row into a table — Excel “Purchase Requests” Table — requester, item, cost, timestamp (fx), Status = “Submitted” (Lab 2).",
        "Condition: if cost is greater than the threshold (e.g. $500) → approval branch; otherwise → auto-approve branch (Lab 3 / Lab 14).",
        "Approval branch: Start and wait for an approval (Approve / Reject); set Assigned to by picking a "
        "real tenant manager from the dropdown — never an external address, or it will hang (Module 2 / 4 reminder).",
        "Approved → Update a row: Status = “Approved”; Send an email (V2) to the requester; notify Procurement.",
        "Rejected → Update a row: Status = “Rejected”; Send an email (V2) to the requester including the approval comments (reason).",
        "Below threshold → Update a row: Status = “Auto-approved”; notify the requester.",
        "Test both branches just above and just below the threshold, plus Approve and Reject; check the "
        "Status column and emails; confirm connections green ✓ (Lab 14).",
    ]),
    ("Question 3", "A3", "Labs 6, 7, 8, 9 — agent components (Module 3)", [
        "Create the agent in Copilot Studio and write its Instructions (e.g. “You are ACME’s customer "
        "assistant. Answer only from the provided knowledge; collect product, quantity and delivery address.”) (Lab 6).",
        "(a) Knowledge / RAG: add ACME’s product sheets, price list and returns/warranty policy as Knowledge "
        "sources (files, or a SharePoint/website). Turn General knowledge OFF so it answers only from your "
        "sources; this is RAG and it prevents hallucination (Lab 7).",
        "Test grounding: ask something in the docs (cited answer) and something not in the docs (it should "
        "say it doesn’t have that information) (Lab 7).",
        "(b) Structured capture: build a Topic with Ask a question nodes storing each answer in a variable "
        "(product, quantity, deliveryAddress); set quantity to Number and re-ask on bad/missing input (Lab 9).",
        "(c) Tool / action: add a tool with a clear name and description — a single connector action (e.g. "
        "Send an email V2) for one step, or a Power Automate agent flow for multi-step work (Lab 8). "
        "With generative orchestration the agent picks the tool from its description, so write it precisely; "
        "if two tools overlap, turn the other one Off while testing.",
        "Refresh the Test pane and confirm the agent answers from knowledge, captures the variables, and calls the tool.",
    ]),
    ("Question 4", "A4", "Labs 10, 11, 15 — end-to-end orchestration (Module 4)", [
        "Capture: the agent collects product, quantity, delivery address into variables (Lab 9 / Lab 15).",
        "Connect agent → flow: build a Power Automate flow whose trigger is “When an agent calls the flow”; "
        "add flow inputs that match the agent variables, and end the flow with “Respond to the agent” "
        "returning a confirmation/status output (Lab 10). It is the same Day-1 flow building — only the trigger changes.",
        "Map the data: in the agent tool, map each variable to its flow input (Fill using → Dynamically "
        "fill with AI, or a Custom value), and show the returned message back in the chat (Lab 10).",
        "Inside the flow: Add a row to the Orders Table (Status = “Received”); use an AI prompt to draft the "
        "order-confirmation text (Lab 11); Condition — if quantity is above the restock threshold, Send an "
        "email (V2) restock alert to the warehouse; Send an email (V2) confirmation to the customer; "
        "Update a row to Status = “Confirmed” (Lab 15).",
        "Every outcome (no dead ends): happy path → logged + confirmed; large order → also raises the restock "
        "alert; bad/missing input → the agent re-asks and the flow validates so it does not crash, then responds with a helpful message.",
        "Test: work through the test log — happy path, threshold just above/below, and bad input; open the "
        "flow’s Run history (every step green); verify the Orders rows and the delivered emails; confirm all "
        "connections green ✓ and recipients are real tenant users (Lab 15 / Lab 16).",
    ]),
]

def case_study_block(d):
    line(d, "C. Case Study", bold=True, size=12, after=4)
    line(d, "Scenario:", bold=True, after=2)
    line(d, SCENARIO, after=4)
    for pt in SCENARIO_POINTS: bullet(d, pt)
    line(d, SCENARIO_TAIL, after=8)

def build_case_paper():
    d = new_doc()
    title_block(d, "Case Study (CS) Assessment — Practical")
    candidate_block(d, 1, [
        "1. This is an individual, open-book practical assessment.",
        "2. A total of 1 hour is given to complete this Case Study assessment.",
        "3. Base every answer on the ACME Pte Ltd case study below.",
        "4. Build the automation in your environment where possible; otherwise describe the exact build steps.",
        "5. All four tasks reuse the techniques you practised in the hands-on labs.",
    ])
    case_study_block(d)
    for label, crit, body in CS_Q:
        runs(d, [(f"{label} ({crit}): ", True), (body, False)], after=4)
        answer_lines(d, 5)
    official_use(d)
    out = os.path.join(OUT, f"Case Study(CS) Assessment - {TITLE}.docx")
    d.save(out); return out

def build_case_answers():
    d = new_doc()
    title_block(d, "Case Study (CS) Assessment — Model Answers (Marking Guide)")
    line(d, "Note to assessor:", bold=True, after=2)
    line(d, "Each task maps to one ability criterion (A1–A4) and to the labs the candidate completed in "
            "class — the model answer is the lab build sequence. Award Competent (C) where the candidate "
            "produces (or correctly describes) a working build that covers the steps in bold; exact "
            "tool clicks and wording will vary.", after=8)
    case_study_block(d)
    for label, crit, lab_ref, steps in CS_A:
        line(d, f"{label} ({crit}) — Model Answer", bold=True, size=12, color=BRAND, after=2)
        runs(d, [("Reference: ", True), (lab_ref, False)], after=4)
        for s in steps: numbered(d, s)
        line(d, "", after=4)
    out = os.path.join(OUT, f"Answers to Case Study(CS) Assessment - {TITLE}.docx")
    d.save(out); return out

if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    for fn in (build_written_paper, build_written_answers, build_case_paper, build_case_answers):
        print("Wrote", fn())
